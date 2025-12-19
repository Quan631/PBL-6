from __future__ import annotations
from pathlib import Path
from typing import List, Dict
import pandas as pd
from docx import Document
from docx.shared import Inches


def export_to_word(
    out_path: Path,
    title: str,
    doc_id: str,
    doc_type: str,
    created_at: str,
    images: List[Dict],
) -> Path:
    """
    images: [{filename, stored_path, ocr_text}]
    """
    doc = Document()
    doc.add_heading(title or f"Document {doc_id}", level=1)
    doc.add_paragraph(f"Document ID: {doc_id}")
    doc.add_paragraph(f"Created at: {created_at}")
    doc.add_paragraph(f"Type: {doc_type}")
    doc.add_paragraph("")

    for idx, img in enumerate(images, start=1):
        doc.add_heading(f"Image {idx}: {img['filename']}", level=2)

        # add image
        try:
            doc.add_picture(str(img["stored_path"]), width=Inches(5.8))
        except Exception:
            doc.add_paragraph("[Could not embed image]")

        doc.add_paragraph("OCR Text:")
        doc.add_paragraph(img.get("ocr_text", "") or "")
        doc.add_page_break()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def export_to_excel(
    out_path: Path,
    doc_id: str,
    doc_type: str,
    created_at: str,
    images: List[Dict],
) -> Path:
    rows = []
    for img in images:
        rows.append({
            "document_id": doc_id,
            "doc_type": doc_type,
            "created_at": created_at,
            "filename": img["filename"],
            "stored_path": str(img["stored_path"]),
            "ocr_text": img.get("ocr_text", "") or "",
        })

    df = pd.DataFrame(rows)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    df.to_excel(out_path, index=False)
    return out_path
